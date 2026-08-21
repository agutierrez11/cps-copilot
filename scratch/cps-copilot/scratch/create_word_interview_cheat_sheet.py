import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

doc = docx.Document()

# Set Margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Primary Colors
PRIMARY_COLOR = RGBColor(15, 23, 42)    # Slate 900
SECONDARY_COLOR = RGBColor(16, 185, 129) # Emerald 500
TEXT_DARK = RGBColor(51, 65, 85)       # Slate 700

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("GUÍA MAESTRA DE PREPARACIÓN: ACI WORLDWIDE")
run.font.name = 'Arial'
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = PRIMARY_COLOR

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle_p.add_run("Antonio Gutiérrez Jiménez | Vacante: BDR LATAM | Entrevistador: Andrés Soler (8:00 AM)")
run_sub.font.name = 'Arial'
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = SECONDARY_COLOR

doc.add_paragraph().paragraph_format.space_after = Pt(10)

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    return p

def add_box(doc, title, text_en, text_es):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    cell.width = Inches(7.0)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    
    r_title = p.add_run(f"❓ {title}\n")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(12)
    r_title.font.bold = True
    r_title.font.color.rgb = PRIMARY_COLOR
    
    r_en_lbl = p.add_run("🇺🇸 EN INGLÉS (B2 Natural):\n")
    r_en_lbl.font.bold = True
    r_en_lbl.font.size = Pt(10)
    r_en_lbl.font.color.rgb = SECONDARY_COLOR
    
    r_en = p.add_run(f"{text_en}\n\n")
    r_en.font.size = Pt(10.5)
    r_en.font.color.rgb = TEXT_DARK
    
    r_es_lbl = p.add_run("🇲🇽 EN ESPAÑOL:\n")
    r_es_lbl.font.bold = True
    r_es_lbl.font.size = Pt(10)
    r_es_lbl.font.color.rgb = RGBColor(71, 85, 105)
    
    r_es = p.add_run(f"{text_es}\n")
    r_es.font.size = Pt(10)
    r_es.font.color.rgb = TEXT_DARK
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Section 1: Pitch Maestro
add_heading(doc, "💎 1. EL ENUNCIADO MAESTRO DE DOBLE DOMINIO (INBOUND + OUTBOUND)")
add_box(
    doc,
    "¿Por qué deberíamos contratarte? (Demostración con Data Real)",
    "\"You should hire me because I dominate BOTH Inbound and Outbound execution with hard data: On Inbound, I converted marketing leads into $16.8M MXN by maximizing response speed; on Outbound, I self-generated $50.5M MXN (72.7% of my $69M portfolio) with deal sizes twice as large as standard inbound leads, leading me to reach 109% quota and the Top 3 National Podium at Clip.\"",
    "\"Deberían contratarme porque domino AMBOS canales, Inbound y Outbound, con datos duros: En Inbound, convertí leads de marketing en $16.8M MXN maximizando la velocidad de respuesta; en Outbound, auto-generé $50.5M MXN (el 72.7% de mi cartera de $69M) con tratos que duplicaron el ticket promedio del canal inbound, lo que me llevó a alcanzar el 109% de cuota y el Podio Nacional Top 3 en Clip.\""
)

# Section 2: Batería STAR con Datos Reales del CV
add_heading(doc, "🎯 2. PREGUNTAS STAR CON DEMOSTRACIÓN DE DATA INBOUND + OUTBOUND")

add_box(
    doc,
    "1. Demuéstrame con DATA cómo manejas Inbound y Outbound (Desglose de Cartera en Clip)",
    "\"At Clip, my $69M MXN total portfolio volume was built on strong execution in both channels:\n• INBOUND DOMINANCE: $16.8M MXN (24.2% of total). I qualified marketing leads instantly, ensuring zero lead decay and high conversion rates from MQL to SQL.\n• OUTBOUND DOMINANCE: $50.5M MXN (72.7% of total). Through active partner prospecting (ERP/POS alliances) and direct C-Level outreach, I generated 15 high-ticket deals with an average size of $555,000 MXN—TWICE the size of inbound leads.\n• TOTAL RESULT: 109% total quota attainment and Top 12% Performer nationally.\"",
    "\"En Clip, mi cartera total de $69M MXN demostró mi dominio en ambos canales:\n• DOMINIO INBOUND: $16.8M MXN (24.2%). Califiqué leads de marketing al instante, garantizando cero enfriamiento de leads y alta conversión de MQL a SQL.\n• DOMINIO OUTBOUND: $50.5M MXN (72.7%). Mediante alianzas con ERPs/POS y contacto directo a C-Levels, generé 15 tratos con un ticket promedio de $555K MXN (el DOBLE que el inbound).\n• RESULTADO: 109% de cumplimiento de cuota y Top 3 del Podio Nacional.\""
)

add_box(
    doc,
    "2. What is your biggest sales achievement? (Historia Retailer Celulares - 200 Sucursales)",
    "\"My biggest achievement was landing a major cell phone retailer with 200 stores.\n• Situation: They had 200 locations with 3 bank terminals per store (600 POS total). They couldn't remove the bank terminals because of credit loan contracts.\n• Task: Close the account without forcing them to break their bank agreements.\n• Action: I proposed adding Clip as a secondary payment channel alongside their bank terminals. I aligned three key stakeholders—Finance, Administration, and IT—to ensure smooth integration.\n• Result: We won the account without breaking their contracts. They scaled to $10 Million MXN in monthly volume, helping me reach the 3rd place overall on Clip's national podium with 109% quota achievement.\"",
    "\"Mi mayor logro fue firmar un retailer de celulares con 200 sucursales (600 TPVs total). No podían quitar las terminales bancarias por contratos de crédito. Mi solución fue agregar Clip como canal secundario sin quitar las otras terminales, sin violar sus contratos. Alineé a Finanzas, Administración y TI, y la cuenta terminó facturando $10M MXN mensuales, llevándome al 3er lugar del podio nacional en Clip con 109% de cuota.\""
)

add_box(
    doc,
    "3. How do you prospect in High-Volume verticals (Gaming, Retail, ERPs)?",
    "\"I combine 3 strategic channels:\n1. Partner Alliances: Channel integration with POS and ERP platforms—like Intelisis, Odoo, Bistrosoft, and Profitroom—gaining direct access to hundreds of warm merchants.\n2. Tech-Enabled Prospecting: I use Sales Navigator, Apollo, and custom scraping workflows to reach C-Level decision-makers.\n3. Deep Pay-In/Pay-Out Diagnostics: In verticals like iGaming or Retail, I focus on approval rates, real-time disbursements, and reducing false positives.\"",
    "\"Combino 3 canales estratégicos: 1) Alianzas con ERPs y POS (Intelisis, Odoo, Bistrosoft, Profitroom) para acceder a cientos de comercios cálidos, 2) Prospección con Sales Navigator y Apollo para contactar decisores C-Level, y 3) Diagnósticos de Pay-In, Pay-Out y reducción de falsos positivos en iGaming y Retail.\""
)

add_box(
    doc,
    "4. How do you summarize the BDR role inside ACI Worldwide?",
    "\"I understand the BDR role at ACI as a hybrid accelerator between Marketing investments and Sales revenue:\n1. Morning (Inbound): Fast lead evaluation (MQL to SQL) to capture intent.\n2. Midday (Outbound Research): AI account research on retail, gaming, and fueling chains in Salesforce/LinkedIn.\n3. Afternoon (Outbound Execution): Targeted phone calls and outreach to book discovery meetings.\n4. End of Day: CRM updates and lead quality feedback to Marketing and Account Executives.\"",
    "\"Entiendo el rol de BDR como un acelerador híbrido entre Marketing y Ventas: 1) Mañana (Inbound): evaluación rápida (MQL a SQL) para capturar el interés, 2) Mediodía (Outbound Research): investigación con IA de cuentas en retail, iGaming y gasolineras, 3) Tarde (Outbound Execution): llamadas enfocadas a C-Levels, 4) Cierre: actualización de CRM y retroalimentación a AEs.\""
)

# Section 3: Insights de Inteligencia ACI
add_heading(doc, "🍊 3. DATOS DE INTELIGENCIA DE ACI WORLDWIDE PARA EXPRIMIR")

table_intel = doc.add_table(rows=1, cols=2)
table_intel.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table_intel.rows[0].cells
set_cell_background(hdr_cells[0], "0F172A")
set_cell_background(hdr_cells[1], "0F172A")

r0 = hdr_cells[0].paragraphs[0].add_run("Dato Extraído de la Empresa")
r0.font.bold = True
r0.font.color.rgb = RGBColor(255, 255, 255)
r1 = hdr_cells[1].paragraphs[0].add_run("Cómo Soltarlo en la Entrevista")
r1.font.bold = True
r1.font.color.rgb = RGBColor(255, 255, 255)

intel_data = [
    ("Vertical Oficial de ACI: Gaming and Digital Entertainment:\nACI ofrece procesamiento de alto volumen, Pay-In/Pay-Out rápido y prevención de fraude para casinos online y entretenimiento digital.",
     "\"Tengo experiencia en consultoría de iGaming y casinos en LATAM. Entiendo que en esta vertical el dolor no es solo cobrar (Pay-In), sino la velocidad de la dispersión (Pay-Out) y evitar que los bloqueos por fraude tiren a jugadores legítimos.\""),
    ("Alberto Olivares (VP LATAM) - El Economista (Junio 2026):\nLos 'falsos positivos' bloquean ventas legítimas y cuestan millones a comercios antes del Mundial 2026.",
     "\"Leí el análisis de Alberto Olivares sobre los falsos positivos: rechazar ventas legítimas por miedo al fraude cuesta millones. La IA de ACI resuelve esto aumentando la tasa de aprobación.\""),
    ("Alberto Olivares (VP LATAM) - El Economista (Abril 2026):\nEl sistema SPEI en México está gravemente subutilizado por los comercios.",
     "\"Coincido con Alberto sobre SPEI: el riel existe pero el comercio en México lo tiene subutilizado. La nueva alianza de ACI con dLocal para ofrecer SPEI y OXXO en 1 sola API nos da el argumento perfecto para calificar SQLs.\""),
    ("Noticia Oficial ACI + dLocal (2026):\nAlianza estratégica para integrar SPEI, OXXO, Pix y Mercado Pago en la plataforma de orquestación ACI.",
     "\"La alianza ACI + dLocal es una ventaja tremenda para prospectar comercios multinacionales que quieren operar en México sin complicaciones de integración.\"")
]

for item, speech in intel_data:
    row_cells = table_intel.add_row().cells
    set_cell_background(row_cells[0], "F8FAFC")
    set_cell_background(row_cells[1], "FFFFFF")
    
    p0 = row_cells[0].paragraphs[0]
    p0.add_run(item).font.size = Pt(9.5)
    
    p1 = row_cells[1].paragraphs[0]
    r_sp = p1.add_run(speech)
    r_sp.font.size = Pt(9.5)
    r_sp.font.italic = True

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Section 4: Tus 3 Preguntas Finales
add_heading(doc, "❓ 4. TUS 3 PREGUNTAS FINAL PARA ANDRÉS SOLER")

p_q = doc.add_paragraph()
p_q.add_run("1. \"In the current LATAM strategy, is the main prospecting focus more on Large Retail/Merchant Orchestration, Gaming, or Financial Institutions?\"\n\n")
p_q.add_run("2. \"What distinguishes a BDR who succeeds quickly and hits quota in their first 90 days at ACI?\"\n\n")
p_q.add_run("3. \"What are the next steps in the interview process after our screening call today?\"")

# Save to multiple exact paths (OneDrive Downloads, System Downloads, Desktop)
target_paths = [
    "C:/Users/Antonio/OneDrive/Downloads/Bateria_Respuestas_ACI_Worldwide_Antonio_Gutierrez_v3.docx",
    "C:/Users/Antonio/OneDrive/Downloads/Bateria_Respuestas_ACI_Worldwide_Antonio_Gutierrez_v2.docx",
    "C:/Users/Antonio/Downloads/Bateria_Respuestas_ACI_Worldwide_Antonio_Gutierrez.docx",
    "C:/Users/Antonio/OneDrive/Desktop/Bateria_Respuestas_ACI_Worldwide_Antonio_Gutierrez.docx",
    "C:/Users/Antonio/Desktop/Bateria_Respuestas_ACI_Worldwide_Antonio_Gutierrez.docx"
]

temp_file = "c:/Users/Antonio/.gemini/antigravity-ide/scratch/cps-copilot/scratch/temp_bateria.docx"
doc.save(temp_file)

saved = []
for path in target_paths:
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        shutil.copy(temp_file, path)
        saved.append(path)
        print(f"[OK] Guardado en: {path}")
    except Exception as e:
        print(f"[SKIP] No se pudo guardar en {path}: {e}")
